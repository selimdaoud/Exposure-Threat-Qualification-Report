#!/usr/bin/env python3
"""Export the executive summary from an HTML report to an editable DOCX file."""

from __future__ import annotations

import argparse
import re
import sys
import tempfile
from pathlib import Path

try:
    from bs4 import BeautifulSoup, Tag
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from PIL import Image, ImageDraw, ImageFont
except ImportError as exc:  # pragma: no cover - exercised when dependencies are absent
    raise SystemExit(
        "Missing dependency. Install the exporter requirements with:\n"
        "  python3 -m pip install beautifulsoup4 python-docx pillow"
    ) from exc


RISK_COLORS = {
    "critical": "B42318",
    "high": "EA580C",
    "moderate": "CA8A04",
    "low": "16A34A",
}
RISK_BACKGROUNDS = {
    "critical": "#FDF4F3",
    "high": "#FFF7ED",
    "moderate": "#FEFCE8",
    "low": "#F0FDF4",
}


def clean_text(node: Tag | None) -> str:
    if node is None:
        return ""
    return re.sub(r"\s+", " ", node.get_text(" ", strip=True)).strip()


def direct_children(node: Tag, name: str) -> list[Tag]:
    return [child for child in node.children if isinstance(child, Tag) and child.name == name]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    if text:
        doc.add_heading(text, level=level)


def add_colored_text(paragraph, text: str, color: str | None = None, bold: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_colored_text(paragraph, text, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def risk_name(node: Tag) -> str:
    classes = node.get("class", [])
    return next((name for name in RISK_COLORS if any(name in css_class for css_class in classes)), "low")


def load_font(size: int, *, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and draw.textbbox((0, 0), candidate, font=font)[2] > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [""]


def diagram_signal_text(text: str) -> str:
    return text.replace("⚡", "").replace("⚑", "").replace("●", "").strip()


def draw_text_lines(draw, xy, lines, font, fill, *, spacing=4):
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += draw.textbbox((x, y), line or " ", font=font)[3] - y + spacing
    return y


def render_heatmap_image(table_html: Tag, output_path: Path) -> None:
    width, margin, gap = 1800, 24, 14
    header_height, card_gap, inner_pad = 58, 12, 12
    column_width = (width - 2 * margin - 3 * gap) // 4
    card_width = column_width - 2 * inner_pad
    title_font = load_font(22, bold=True)
    owner_font = load_font(19, bold=True)
    body_font = load_font(17)
    small_font = load_font(15)
    headers = table_html.select("thead th")
    cells = table_html.select("tbody > tr > td")

    column_cards: list[list[dict]] = []
    for html_cell in cells:
        cards = []
        for card in html_cell.select(".hm-owner-card"):
            cards.append({
                "risk": risk_name(card),
                "score": clean_text(card.select_one(".hm-card-badge")),
                "owner": clean_text(card.select_one(".hm-owner-name")),
                "host": clean_text(card.select_one(".hm-host-name")),
                "count": clean_text(card.select_one(".hm-cve-count")),
                "eol": clean_text(card.select_one(".hm-eol-badge")),
                "signals": diagram_signal_text(clean_text(card.select_one(".hm-signals"))),
            })
        column_cards.append(cards)

    measure = ImageDraw.Draw(Image.new("RGB", (width, 100), "white"))
    column_heights = []
    for cards in column_cards:
        total = header_height + inner_pad
        for card in cards:
            lines = (
                wrap_text(measure, card["owner"], owner_font, card_width - 92)
                + wrap_text(measure, card["host"], body_font, card_width - 24)
                + [card["count"]]
            )
            if card["eol"]:
                lines += wrap_text(measure, card["eol"], small_font, card_width - 24)
            if card["signals"]:
                lines += wrap_text(measure, card["signals"], small_font, card_width - 24)
            total += 28 + sum(
                measure.textbbox((0, 0), line or " ", font=owner_font if index == 0 else body_font)[3] + 4
                for index, line in enumerate(lines)
            ) + card_gap
        column_heights.append(total + inner_pad)
    height = max(max(column_heights, default=0), 240)

    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    for index, (header, html_cell) in enumerate(zip(headers, cells)):
        risk = risk_name(header)
        x0 = margin + index * (column_width + gap)
        x1 = x0 + column_width
        draw.rounded_rectangle((x0, 0, x1, height - 1), radius=12, fill=RISK_BACKGROUNDS[risk], outline="#CBD5E1", width=2)
        label = clean_text(header).upper()
        label_box = draw.textbbox((0, 0), label, font=title_font)
        draw.text((x0 + (column_width - label_box[2]) / 2, 16), label, font=title_font, fill="#" + RISK_COLORS[risk])
        draw.line((x0, header_height, x1, header_height), fill="#" + RISK_COLORS[risk], width=5)
        y = header_height + inner_pad
        cards = column_cards[index]
        if not cards:
            draw.text((x0 + inner_pad, y), "-", font=body_font, fill="#64748B")
        for card in cards:
            owner_lines = wrap_text(draw, card["owner"], owner_font, card_width - 92)
            other_lines = wrap_text(draw, card["host"], body_font, card_width - 24) + [card["count"]]
            if card["eol"]:
                other_lines += wrap_text(draw, card["eol"], small_font, card_width - 24)
            if card["signals"]:
                other_lines += wrap_text(draw, card["signals"], small_font, card_width - 24)
            card_height = 28
            card_height += sum(draw.textbbox((0, 0), line or " ", font=owner_font)[3] + 4 for line in owner_lines)
            card_height += sum(draw.textbbox((0, 0), line or " ", font=body_font)[3] + 4 for line in other_lines)
            draw.rounded_rectangle((x0 + inner_pad, y, x1 - inner_pad, y + card_height), radius=10, fill="#FFFFFF", outline="#" + RISK_COLORS[card["risk"]], width=2)
            score = card["score"]
            badge_box = draw.textbbox((0, 0), score, font=small_font)
            badge_width = max(42, badge_box[2] + 18)
            draw.rounded_rectangle((x1 - inner_pad - badge_width - 8, y + 9, x1 - inner_pad - 8, y + 38), radius=14, fill="#" + RISK_COLORS[card["risk"]])
            draw.text((x1 - inner_pad - badge_width / 2 - badge_box[2] / 2 - 8, y + 15), score, font=small_font, fill="#FFFFFF")
            text_y = draw_text_lines(draw, (x0 + inner_pad + 12, y + 12), owner_lines, owner_font, "#0F172A")
            draw.line((x0 + inner_pad + 12, text_y + 1, x1 - inner_pad - 12, text_y + 1), fill="#E2E8F0", width=1)
            text_y += 8
            draw_text_lines(draw, (x0 + inner_pad + 12, text_y), other_lines, body_font, "#475569")
            y += card_height + card_gap
    image.save(output_path)


def add_overview(doc: Document, soup: BeautifulSoup, executive_summary: Tag) -> None:
    title = clean_text(soup.select_one(".topbar h1")) or "Exposure & Threat Qualification Report"
    doc.add_heading(title, 0)
    for meta in soup.select(".topbar .meta"):
        doc.add_paragraph(clean_text(meta))

    add_heading(doc, "Executive Summary")
    summary_bar = executive_summary.find("summary", recursive=False)
    if summary_bar is None:
        return

    scope = clean_text(summary_bar.select_one(".exec-scope-line"))
    headline = clean_text(summary_bar.select_one(".exec-headline"))
    posture = clean_text(summary_bar.select_one(".posture-badge"))
    pills = [clean_text(pill) for pill in summary_bar.select(".exec-summary-pills .badge")]

    if scope:
        doc.add_paragraph(scope)
    if headline:
        doc.add_paragraph(headline)
    if posture:
        paragraph = doc.add_paragraph()
        add_colored_text(paragraph, posture, bold=True)
    if pills:
        doc.add_paragraph("Risk exposure: " + " | ".join(pills))


def add_heatmap(doc: Document, executive_summary: Tag) -> None:
    table_html = executive_summary.select_one(".heatmap-table")
    if table_html is None:
        return

    add_heading(doc, "Risk Exposure Matrix - Owner x Risk Level", level=2)
    with tempfile.TemporaryDirectory() as temp_dir:
        image_path = Path(temp_dir) / "risk-exposure-matrix.png"
        render_heatmap_image(table_html, image_path)
        doc.add_picture(str(image_path), width=Inches(10.45))


def add_business_drivers(doc: Document, executive_summary: Tag) -> None:
    table_html = executive_summary.select_one(".kbd-table")
    if table_html is None:
        return

    add_heading(doc, "Key Business Drivers", level=2)
    headers = [clean_text(cell) for cell in table_html.select("thead th")]
    body_rows = table_html.select("tbody > tr")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.cell(0, index), header, bold=True)

    score_details: list[tuple[str, str]] = []
    for html_row in body_rows:
        cells = direct_children(html_row, "td")
        row = table.add_row()
        for index, html_cell in enumerate(cells):
            if index == len(headers) - 1:
                score_badge = html_cell.select_one(".kbd-score-badge")
                value = clean_text(score_badge) or clean_text(html_cell)
                popup = html_cell.select_one(".kbd-score-popup")
                if popup is not None:
                    label = " | ".join(clean_text(cell) for cell in cells[:5])
                    score_details.append((label, clean_text(popup)))
            else:
                value = clean_text(html_cell)
            set_cell_text(row.cells[index], value)

    if score_details:
        add_heading(doc, "Host Score Breakdowns", level=3)
        for label, detail in score_details:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_colored_text(paragraph, label + ": ", bold=True)
            paragraph.add_run(detail)


def add_technical_drivers(doc: Document, executive_summary: Tag) -> None:
    drivers = executive_summary.select_one(".exec-drivers")
    if drivers is None:
        return

    add_heading(doc, "Key Technical Drivers", level=2)
    columns = drivers.select(":scope > .exec-drivers-col")
    if columns:
        for row in columns[0].select(":scope > .exec-driver-row"):
            paragraph = doc.add_paragraph(style="List Bullet")
            expandable = row.select_one(".driver-expand")
            groups = row.select(".driver-owner-expand")
            if expandable is not None:
                expandable.extract()
            paragraph.add_run(clean_text(row))
            for group in groups:
                owner_summary = group.find("summary", recursive=False)
                if owner_summary is not None:
                    doc.add_paragraph(clean_text(owner_summary), style="List Bullet 2")
                for machine in group.select(":scope > .driver-machine-expand"):
                    machine_summary = machine.find("summary", recursive=False)
                    if machine_summary is not None:
                        doc.add_paragraph(clean_text(machine_summary), style="List Bullet 3")
    if len(columns) > 1:
        add_heading(doc, "Required Actions", level=3)
        for row in columns[1].select(":scope > .exec-action-row"):
            doc.add_paragraph(clean_text(row), style="List Bullet")


def add_methodology(doc: Document, executive_summary: Tag) -> None:
    method = executive_summary.select_one(".posture-method-body")
    if method is None:
        return

    add_heading(doc, "How is the Risk Posture determined?", level=2)
    for child in method.children:
        if not isinstance(child, Tag):
            continue
        if child.name == "table":
            rows = child.select(":scope > tr")
            if not rows:
                rows = child.select("tr")
            width = max((len(direct_children(row, "td")) for row in rows), default=0)
            if not width:
                continue
            table = doc.add_table(rows=0, cols=width)
            table.style = "Table Grid"
            for html_row in rows:
                row = table.add_row()
                for index, html_cell in enumerate(direct_children(html_row, "td")):
                    set_cell_text(row.cells[index], clean_text(html_cell))
        else:
            text = clean_text(child)
            if text:
                doc.add_paragraph(text)


def add_summary_points(doc: Document, executive_summary: Tag) -> None:
    points = executive_summary.select(".exec-summary-list > li")
    if points:
        add_heading(doc, "Summary Points", level=2)
        for point in points:
            doc.add_paragraph(clean_text(point), style="List Bullet")
    footer = clean_text(executive_summary.select_one(".exec-summary-footer"))
    if footer:
        doc.add_paragraph(footer)


def _load_template(template_path: Path) -> Document:
    import copy
    from docx.oxml.ns import qn
    from docx.styles import BabelFish
    doc = Document(str(template_path))
    styles_el = doc.styles.element
    seen: set[str] = set()
    for style in list(styles_el.findall(qn("w:style"))):
        sid = style.get(qn("w:styleId"), "")
        if sid in seen:
            # python-docx's style lookup breaks on duplicate styleId entries; remove them.
            styles_el.remove(style)
            continue
        seen.add(sid)
        # Normalize w:name values: python-docx looks up by BabelFish.ui2internal() names
        # (e.g. "heading 1"), but some templates store the UI name ("Heading 1") instead.
        name_el = style.find(qn("w:name"))
        if name_el is not None:
            current = name_el.get(qn("w:val"), "")
            normalised = BabelFish.ui2internal(current)
            if normalised != current:
                name_el.set(qn("w:val"), normalised)
    # Merge any styles missing from the template from a default Document so the template
    # doesn't need to define every built-in style the exporter uses (Table Grid, List Bullet…).
    # Template styles take precedence; only missing ones are added.
    default_styles_el = Document().styles.element
    for default_style in default_styles_el.findall(qn("w:style")):
        sid = default_style.get(qn("w:styleId"), "")
        if sid not in seen:
            styles_el.append(copy.deepcopy(default_style))
    return doc


def export_docx(input_path: Path, output_path: Path, template_path: Path | None = None) -> None:
    soup = BeautifulSoup(input_path.read_text(encoding="utf-8"), "html.parser")
    executive_summary = soup.select_one(".exec-summary")
    if executive_summary is None:
        raise ValueError(f"No executive summary found in {input_path}")

    document = _load_template(template_path) if template_path else Document()
    if not template_path:
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        document.styles["Normal"].font.name = "Aptos"
        document.styles["Normal"].font.size = Pt(9)

    add_overview(document, soup, executive_summary)
    add_heatmap(document, executive_summary)
    add_business_drivers(document, executive_summary)
    add_technical_drivers(document, executive_summary)
    add_methodology(document, executive_summary)
    add_summary_points(document, executive_summary)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export the executive summary from an oracle-cve-intel HTML report to DOCX."
    )
    parser.add_argument("html_report", type=Path, help="Path to the generated HTML report")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="DOCX output path (default: HTML report path with a .docx suffix)",
    )
    parser.add_argument(
        "-t",
        "--template",
        type=Path,
        help="Path to a DOCX template file whose styles and page layout will be used",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    output_path = args.output or args.html_report.with_suffix(".docx")
    try:
        export_docx(args.html_report, output_path, template_path=args.template)
    except (OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Executive summary written to {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
